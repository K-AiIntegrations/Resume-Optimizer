from typing import Tuple
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import os, time

from ..models import Resume

def export_docx(resume: Resume, out_path: str) -> str:
    doc = Document()
    styles = doc.styles
    for s in ["Normal"]:
        styles[s].font.name = "Calibri"
        styles[s].font.size = Pt(11)

    # Summary
    if resume.pii:
        p = doc.add_paragraph()
        p.add_run(resume.pii.get("name",""))
        p.add_run(" | "+(resume.pii.get("email","")))

    if resume.summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(resume.summary)

    # Experience (raw)
    doc.add_heading("Experience", level=1)
    doc.add_paragraph(resume.raw_text[:2000])  # MVP

    doc.save(out_path)
    return out_path

def export_pdf(resume: Resume, out_path: str) -> str:
    c = canvas.Canvas(out_path, pagesize=letter)
    width, height = letter
    textobject = c.beginText(40, height-50)
    textobject.setFont("Times-Roman", 11)
    body = (resume.summary + "\n\n" + resume.raw_text)[:6000]
    for line in body.split("\n"):
        textobject.textLine(line[:100])
    c.drawText(textobject)
    c.showPage()
    c.save()
    return out_path

def export_resume_files(resume: Resume, base_dir: str) -> Tuple[str,str]:
    ts = int(time.time())
    docx_path = os.path.join(base_dir, f"resume_{ts}.docx")
    pdf_path = os.path.join(base_dir, f"resume_{ts}.pdf")
    export_docx(resume, docx_path)
    export_pdf(resume, pdf_path)
    return docx_path, pdf_path
